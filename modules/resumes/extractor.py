'''
Author:     RapidApply Contributors
LinkedIn:   

Copyright (C) 2024 RapidApply Contributors

License:    GNU Affero General Public License
            https://www.gnu.org/licenses/agpl-3.0.en.html
            
GitHub:     https://github.com/Nickalus12/RapidApply

'''

import os
from typing import Optional
from modules.helpers import print_lg

def extract_text_from_pdf(pdf_path: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    try:
        import PyPDF2
        
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        
        return text.strip()
        
    except ImportError:
        print_lg("PyPDF2 not installed. Trying alternative method...")
        # Try pdfplumber as fallback
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
            
        except ImportError:
            print_lg("No PDF library available. Install PyPDF2 or pdfplumber.")
            # Return filename as fallback
            return f"[PDF Content from {os.path.basename(pdf_path)}]"
    
    except Exception as e:
        print_lg(f"Error extracting PDF text: {str(e)}")
        return f"[Error reading PDF: {os.path.basename(pdf_path)}]"

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
        
        doc = Document(docx_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text.strip()
        
    except ImportError:
        print_lg("python-docx not installed. Please install it to read DOCX files.")
        # Return filename as fallback
        return f"[DOCX Content from {os.path.basename(docx_path)}]"
    
    except Exception as e:
        print_lg(f"Error extracting DOCX text: {str(e)}")
        return f"[Error reading DOCX: {os.path.basename(docx_path)}]"

# Legacy imports for backward compatibility
from config.personals import *
from config.questions import default_resume_path



